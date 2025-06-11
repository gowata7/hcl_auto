import markdown
import docx
import re  
import pandas as pd
import pypandoc

from bs4 import BeautifulSoup
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.oxml.ns import qn, nsdecls
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

document = Document()
sections = document.sections

for section in sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    
section = document.sections[0]

sectPr = section._sectPr
cols = sectPr.xpath('./w:cols')[0]
cols.set(qn('w:num'), '1')
    
# pypandoc.pandoc_path = 'C:\\Users\\Administrator\\AppData\\Local\\Programs\\Python\\Python310\\Lib\\site-packages\\pypandoc'
# pandoc_path = 'C:\\Users\\Administrator\\AppData\\Local\\Programs\\Python\\Python310\\Lib\\site-packages\\pypandoc'
# os.environ['PATH'] += os.pathsep + pandoc_path

# set fixed column width size (test not yet)
# def set_col_widths(table):
#     widths = (Inches(1), Inches(2), Inches(1.5))
#     for row in table.rows:
#         for idx, width in enumerate(widths):
#             row.cells[idx].width = width

# def set_autofit(table):
    # for idx
    
def set_cell_font(cell, font_size=10):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(font_size)
    
def handle_tag_to_rtf(cell, tag):
    """
    Convert HTML content to rich text format (RTF) string.
    """
    rtf_string = ""  # Initialize the RTF string
    # print("HTML 데이터:", tag)
    
    # Handle NavigableString (text content)
    if isinstance(tag, str):
        # Replace <br/> tags with newline characters
        tag = tag.replace("<br/>", "\n")
        return tag.strip()
    else:
        # Handle Tag (HTML element)
        for content in tag.contents:
            if isinstance(content, str):
                # Replace <br/> tags with newline characters
                content = content.replace("<br/>", "\n")
                rtf_string += content.strip()
            else:
                # Handle nested tags recursively
                if content.name == 'b':
                    # Handle <b> (bold) tag
                    rtf_string += f"<b>{handle_tag_to_rtf(cell, content)}</b>"
                elif content.name == 'i':
                    # Handle <i> (italic) tag
                    rtf_string += f"<i>{handle_tag_to_rtf(cell, content)}</i>"
                else:
                    # Handle other tags recursively
                    _rtf_string = handle_tag_to_rtf(cell, content)
                    if _rtf_string:
                        rtf_string += _rtf_string
   
    # Add accumulated content to the cell
    cell.add_paragraph(rtf_string.strip())


               
def handle_tag_to_nested_table(cell, tag):
    """
    Handle rich text with HTML tags in a cell.
   
    Args:
    - cell: The cell to which the rich text should be added.
    - tag: The HTML tag containing the nested table.
    """
    if isinstance(tag, str) and "<" in tag and ">" in tag:
        soup = BeautifulSoup(tag, "html.parser")
        table_elements = soup.find_all("table")
        if table_elements:
            nested_table_html = table_elements[0]          
            nested_table = cell.add_table(rows=0, cols=len(nested_table_html.find_all('tr')[0].find_all(['td', 'th'])))
            nested_table.allow_autofit = False
           
            for row in nested_table_html.find_all('tr'):
                nested_row = nested_table.add_row().cells
                for idx, cell_html in enumerate(row.find_all(['td', 'th'])):
                    cell_text = cell_html.get_text().strip()
                    nested_row[idx].text = cell_text
    else:
        cell.text = str(tag).strip()  # Set the cell text directly if it's not a table
    return cell.text


def convert_markdown_to_rtf(cell, val):
    # Convert Markdown to HTML
    rtf_content = pypandoc.convert_text(val, 'plain', format='html')
    # cell.text = convert_markdown_to_rtf(rtf_content)
    cell.text = rtf_content
    return cell.text


def addSummaryTable_all(df):
    
#     if [ month != '누적' ]:
#         df = df[df['월'] == month].reset_index()
#     else:
#         pass
    
#     print("doc", df)
    
    df.sort_values(by=['진행상태', '리전'], axis=0, ascending=True, inplace=True)
    table = document.add_table(rows=1, cols=4, style="Table Grid")

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '진행상태'
    hdr_cells[1].text = '리전'
    hdr_cells[2].text = '센터'
    hdr_cells[3].text = '내용 상세'

    hdr_cells[0].width = Cm(0.05)
    hdr_cells[1].width = Cm(0.05)
    hdr_cells[2].width = Cm(0.05)
    hdr_cells[3].width = Cm(50)
    
    for ind in df.index:
        row_cells = table.add_row().cells
        row_cells[0].text = df['진행상태'][ind]
        row_cells[1].text = df['리전'][ind]
        row_cells[2].text = str(df['AZ'][ind])
        # row_cells[3].text = df['작업내용'][ind]  
        
        cell_content = df['작업내용'][ind]
        convert_markdown_to_rtf(row_cells[3], cell_content)
        
        for cell in row_cells:
            set_cell_font(cell, font_size=10)

def addSummaryTable(df):
    
    # if [ month != '누적' ]:
    #     df = df[df['월'] == month].reset_index()
    # else:
    #     pass
    
    df.sort_values(by=['진행상태','리전'], axis=0, ascending=True, inplace=True)
    table = document.add_table(rows=1, cols=2, style="Table Grid")

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '진행상태'
    hdr_cells[1].text = '내용 상세'

    for ind in df.index:
        row_cells = table.add_row().cells
        row_cells[0].text = df['진행상태'][ind]
        
        # detailed = df['작업내용'][ind]
        cell_content = df['작업내용'][ind]
        convert_markdown_to_rtf(row_cells[1], cell_content)
        # row_cells[1].text = df['작업내용'][ind]
        for cell in row_cells:
            set_cell_font(cell, font_size=10)
        
def addTable3(df):
    df.sort_values(by=['availability_zone'], axis=0, ascending=False, inplace=True)
    table = document.add_table(rows=1, cols=5, style="Table Grid")
    hdr_cells = table.rows[0].cells
    headers = ['Name', 'Status', 'Flavor', 'Address', 'Availability_zone']
    for i, header in enumerate(headers):
        paragraph = hdr_cells[i].paragraphs[0]
        run = paragraph.add_run(header)
        run.bold = True

    for ind in df.index:
        row_cells = table.add_row().cells
        row_cells[0].text = df['name'][ind] if pd.notna(df['name'][ind]) else "None"
        row_cells[1].text = df['status'][ind] if pd.notna(df['status'][ind]) else "None"
        row_cells[2].text = df['flavor'][ind] if pd.notna(df['flavor'][ind]) else "None"
        row_cells[3].text = df['address'][ind] if pd.notna(df['address'][ind]) else "None"
        row_cells[4].text = df['availability_zone'][ind] if pd.notna(df['availability_zone'][ind]) else "None"

        for cell in row_cells:
            set_cell_font(cell, font_size=10)

def addTable_instance(df):
    # df.sort_values(by=['availability_zone'], axis=0, ascending=False, inplace=True)
    table = document.add_table(rows=1, cols=7, style="Table Grid")
    hdr_cells = table.rows[0].cells
    headers = ['Date', 'Project', 'Name', 'Flavor', 'Address', 'AZ', 'Host']
    for i, header in enumerate(headers):
        paragraph = hdr_cells[i].paragraphs[0]
        run = paragraph.add_run(header)
        run.bold = True

    for ind in df.index:
        row_cells = table.add_row().cells
        row_cells[0].text = df['date'][ind].strftime('%Y-%m-%d') if pd.notna(df['date'][ind]) else "None"
        row_cells[1].text = df['project_id'][ind] if pd.notna(df['project_id'][ind]) else "None"
        row_cells[2].text = df['instance_name'][ind] if pd.notna(df['instance_name'][ind]) else "None"
        row_cells[3].text = df['flavor'][ind] if pd.notna(df['flavor'][ind]) else "None"
        row_cells[4].text = df['address'][ind] if pd.notna(df['address'][ind]) else "None"
        row_cells[5].text = df['availability_zone'][ind] if pd.notna(df['availability_zone'][ind]) else "None"
        row_cells[6].text = df['hostname'][ind] if pd.notna(df['hostname'][ind]) else "None"

        for cell in row_cells:
            set_cell_font(cell, font_size=10)

def addTable_loadbalancer(df):
    # df.sort_values(by=['availability_zone'], axis=0, ascending=False, inplace=True)
    table = document.add_table(rows=1, cols=6, style="Table Grid")
    hdr_cells = table.rows[0].cells
    headers = ['Date', 'Tenant', 'Name', 'VIP', 'Port', 'Protocol']
    for i, header in enumerate(headers):
        paragraph = hdr_cells[i].paragraphs[0]
        run = paragraph.add_run(header)
        run.bold = True

    for ind in df.index:
        row_cells = table.add_row().cells
        row_cells[0].text = df['date'][ind].strftime('%Y-%m-%d') if pd.notna(df['date'][ind]) else "None"
        row_cells[1].text = df['tenant'][ind] if pd.notna(df['tenant'][ind]) else "None"
        row_cells[2].text = df['lb_name'][ind] if pd.notna(df['lb_name'][ind]) else "None"
        row_cells[3].text = df['lb_vip'][ind] if pd.notna(df['lb_vip'][ind]) else "None"
        row_cells[4].text = df['port'][ind] if pd.notna(df['port'][ind]) else "None"
        row_cells[5].text = df['service_type'][ind] if pd.notna(df['service_type'][ind]) else "None"

        for cell in row_cells:
            set_cell_font(cell, font_size=10)

def addTable_LB(df):
    df.sort_values(by=['Tenant'], axis=0, ascending=False, inplace=True)
    table = document.add_table(rows=1, cols=5, style="Table Grid")
    hdr_cells = table.rows[0].cells
    headers = ['Tenant', 'LB_Name', 'LB_VIP', 'Port', 'Service_Type']
    for i, header in enumerate(headers):
        paragraph = hdr_cells[i].paragraphs[0]
        run = paragraph.add_run(header)
        run.bold = True

    for ind in df.index:
        row_cells = table.add_row().cells
        row_cells[0].text = df['Tenant'][ind] if pd.notna(df['Tenant'][ind]) else "None"
        row_cells[1].text = df['LB_Name'][ind] if pd.notna(df['LB_Name'][ind]) else "None"
        row_cells[2].text = df['LB_VIP'][ind] if pd.notna(df['LB_VIP'][ind]) else "None"
        row_cells[3].text = str(df['Port'][ind]) if pd.notna(df['Port'][ind]) else "None"
        row_cells[4].text = df['Service_Type'][ind] if pd.notna(df['Service_Type'][ind]) else "None"

        for cell in row_cells:
            set_cell_font(cell, font_size=10)
            
def addTable3_all(df):
    
    # if [ month != '누적' ]:
    #     df = df[df['월'] == month].reset_index()
    # else:
    #     pass
    
    df.sort_values(by=['작업유형','리전'], axis=0, ascending=False, inplace=True)
    table = document.add_table(rows=1, cols=5, style="Table Grid")
    table.allow_autofit = True
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '작업유형'
    hdr_cells[1].text = '리전'
    hdr_cells[2].text = 'AZ'
    hdr_cells[3].text = '진행상태'
    hdr_cells[4].text = '제목'

    hdr_cells[0].width = Cm(0.1)
    hdr_cells[1].width = Cm(0.05)
    hdr_cells[2].width = Cm(0.05)
    hdr_cells[3].width = Cm(0.05)
    hdr_cells[4].width = Cm(50)
    
    for ind in df.index:
        row_cells = table.add_row().cells
        row_cells[0].text = df['작업유형'][ind] 
        row_cells[1].text = df['리전'][ind]
        row_cells[2].text = str(df['AZ'][ind])
        row_cells[3].text = df['진행상태'][ind] 
        row_cells[4].text = df['제목'][ind] 
        
        for cell in row_cells:
            set_cell_font(cell, font_size=10)
        
def addTable_capacity_all(df,  month):
    
    if [ month != '누적' ]:
        df = df[df['월'] == month].reset_index()
    else:
        pass
    
    df.sort_values(by=['진행상태','카테고리','리전'], axis=0, ascending=True, inplace=True)
    table = document.add_table(rows=1, cols=5, style="Table Grid")
    table.allow_autofit = True
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '리전'
    hdr_cells[1].text = 'AZ'
    hdr_cells[2].text = '진행상태'
    hdr_cells[3].text = '카테고리'
    hdr_cells[4].text = '제목'

    hdr_cells[0].width = Cm(0.05)
    hdr_cells[1].width = Cm(0.05)
    hdr_cells[2].width = Cm(0.05)
    hdr_cells[3].width = Cm(0.05)
    hdr_cells[4].width = Cm(50)
    
    for ind in df.index:
        row_cells = table.add_row().cells
        row_cells[0].text = df['리전'][ind]
        row_cells[1].text = str(df['AZ'][ind])
        row_cells[2].text = df['진행상태'][ind] 
        row_cells[3].text = df['카테고리'][ind] 
        row_cells[4].text = df['제목'][ind] 
        for cell in row_cells:
            set_cell_font(cell, font_size=10)
            
def addTable_capacity(df,  month):
    
    if [ month != '누적' ]:
        df = df[df['월'] == month].reset_index()
    else:
        pass
    
    df.sort_values(by=['카테고리'], axis=0, ascending=False, inplace=True)
    table = document.add_table(rows=1, cols=3, style="Table Grid")
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '진행상태'
    hdr_cells[1].text = '카테고리'
    hdr_cells[2].text = '제목'

    for ind in df.index:
        row_cells = table.add_row().cells
        row_cells[0].text = df['진행상태'][ind]
        row_cells[1].text = df['카테고리'][ind]         
        row_cells[2].text = df['제목'][ind]
        for cell in row_cells:
            set_cell_font(cell, font_size=10)
        
def addTable_security(df,  month):
    
    if [ month != '누적' ]:
        df = df[df['월'] == month].reset_index()
    else:
        pass
    
    df.sort_values(by=['진행상태'], axis=0, ascending=False, inplace=True)
    table = document.add_table(rows=1, cols=2, style="Table Grid")
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '진행상태'
    hdr_cells[1].text = '제목'

    for ind in df.index:
        row_cells = table.add_row().cells
        row_cells[0].text = df['진행상태'][ind]
        row_cells[1].text = df['제목'][ind] 
        for cell in row_cells:
            set_cell_font(cell, font_size=10)
        
def addTable_security_all(df,  month):
    
    if [ month != '누적' ]:
        df = df[df['월'] == month].reset_index()
    else:
        pass
    
    df.sort_values(by=['리전','진행상태'], axis=0, ascending=False, inplace=True)
    table = document.add_table(rows=1, cols=3, style="Table Grid")
    table.allow_autofit = True
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '리전'
    hdr_cells[1].text = '진행상태'
    hdr_cells[2].text = '제목'

    hdr_cells[0].width = Cm(0.05)
    hdr_cells[1].width = Cm(0.05)
    hdr_cells[2].width = Cm(50)
    
    for ind in df.index:
        row_cells = table.add_row().cells
        row_cells[0].text = df['리전'][ind]
        row_cells[1].text = df['진행상태'][ind] 
        row_cells[2].text = df['제목'][ind]         
        for cell in row_cells:
            set_cell_font(cell, font_size=10)
            
def addTable_insp(df,  month):
    
    if [ month != '누적' ]:
        df = df[df['월'] == month].reset_index()
    else:
        pass
    
    df.sort_values(by=['진행상태'], axis=0, ascending=False, inplace=True)
    table = document.add_table(rows=1, cols=3, style="Table Grid")
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '진행상태'
    hdr_cells[1].text = '벤더'
    hdr_cells[2].text = '상세 현황'

    for ind in df.index:
        row_cells = table.add_row().cells
        row_cells[0].text = df['진행상태'][ind]
        row_cells[1].text = df['공급사'][ind] 
        row_cells[2].text = df['제목'][ind] 
        for cell in row_cells:
            set_cell_font(cell, font_size=10)
            
def addTable_insp_all(df,  month):
    
    if [ month != '누적' ]:
        df = df[df['월'] == month].reset_index()
    else:
        pass
    
    df.sort_values(by=['진행상태','리전'], axis=0, ascending=False, inplace=True)
    table = document.add_table(rows=1, cols=4, style="Table Grid")
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '리전'
    hdr_cells[1].text = '진행상태'
    hdr_cells[2].text = '벤더'
    hdr_cells[3].text = '비고'

    hdr_cells[0].width = Cm(0.05)
    hdr_cells[1].width = Cm(0.05)
    hdr_cells[2].width = Cm(0.05)
    hdr_cells[3].width = Cm(50)
    
    for ind in df.index:
        row_cells = table.add_row().cells
        row_cells[0].text = df['리전'][ind]
        row_cells[1].text = df['진행상태'][ind]
        row_cells[2].text = df['공급사'][ind] 
        row_cells[3].text = df['제목'][ind] 
        for cell in row_cells:
            set_cell_font(cell, font_size=10)
        
def addTablek8s(df_filtered,  month):
    
    # if [ month != '누적' ]:
    #     df = df[df['월'] == month].reset_index()
    # else:
    #     pass
    
    df_filtered = df_filtered[[ '리전','클러스터명', '프로메테우스 설치 유무']]
    # df_filtered.sort_values(by=['클러스터명'], axis=0, ascending=False, inplace=True)
    df_filtered.sort_values(by=['리전', '클러스터명'], axis=0, ascending=False, inplace=True)
    unique_df = df_filtered['클러스터명'].value_counts()
    unique_df = pd.DataFrame({'클러스터명':unique_df.index, '노드수':unique_df.values}) 
    # unique_df.sort_values(by=['클러스터명'], axis=0, ascending=False, inplace=True)
    unique_df.sort_values(by=['클러스터명'], axis=0, ascending=False, inplace=True)
    unique_df = unique_df.reset_index()
    df_filtered = df_filtered.drop_duplicates(['클러스터명'])
    
    table = document.add_table(rows=1, cols=4, style="Table Grid")
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '리전'
    hdr_cells[1].text = '클러스터명'
    hdr_cells[2].text = '노드수'
    hdr_cells[3].text = '프로메테우스 설치 유무'
    
    print(unique_df)
    print(type(unique_df))
    print(df_filtered)
    print(df_filtered.index)
    print(unique_df.index)
    
    i = 0
    for ind in df_filtered.index:
        row_cells = table.add_row().cells
        row_cells[0].text = df_filtered['리전'][ind]
        row_cells[1].text = df_filtered['클러스터명'][ind]
        row_cells[2].text = str(unique_df['노드수'][i])
        row_cells[3].text = df_filtered['프로메테우스 설치 유무'][ind] 
        i += 1
        for cell in row_cells:
            set_cell_font(cell, font_size=10)
            
def addTablek8s_all(df_filtered,  month):
    
    # if [ month != '누적' ]:
    #     df = df[df['월'] == month].reset_index()
    # else:
    #     pass
    # mask = df_filtered['node_type'].isin(['local','Rancher'])
    # df_filtered = df_filtered[~mask]
        # df_filtered = df_filtered[df_filtered.columns.difference(['node_type'])]
        # df_filtered = df_filtered[df_filtered['node_type'] != 'local']
        # print(df_filtered)
    df_filtered = df_filtered[[ '리전','클러스터명', '프로메테우스 설치 유무']]
    df_filtered.sort_values(by=['클러스터명'], axis=0, ascending=False, inplace=True)
    # df_filtered.sort_values(by=['리전','node_type'], axis=0, ascending=False, inplace=True)
    print(df_filtered)
    
    unique_df = df_filtered['클러스터명'].value_counts()
    unique_df = pd.DataFrame({'클러스터명':unique_df.index, '노드수':unique_df.values}) 
    unique_df.sort_values(by=['클러스터명'], axis=0, ascending=False, inplace=True)
    unique_df = unique_df.reset_index()
    df_filtered = df_filtered.drop_duplicates(['클러스터명'])

    table = document.add_table(rows=1, cols=4, style="Table Grid")
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '리전'
    hdr_cells[1].text = '클러스터명'
    hdr_cells[2].text = '노드수'
    hdr_cells[3].text = '프로메테우스 설치 유무'
    
    print(unique_df)
    # print(type(unique_df))
    # print(df_filtered)
    # print(df_filtered.index)
    print(unique_df.index)
    
    i = 0
    for ind in df_filtered.index:
        row_cells = table.add_row().cells
        row_cells[0].text = df_filtered['리전'][ind]
        row_cells[1].text = df_filtered['클러스터명'][ind]
        row_cells[2].text = str(unique_df['노드수'][i])
        row_cells[3].text = df_filtered['프로메테우스 설치 유무'][ind] 
        i += 1
        for cell in row_cells:
            set_cell_font(cell, font_size=10)
        
def addTable4_all(df):
    
    # if [ month != '누적' ]:
    #     df = df[df['월'] == month].reset_index()
    # else:
    #     pass
    
    df.sort_values(by=['진행상태','리전'], axis=0, ascending=True, inplace=True)
    table = document.add_table(rows=1, cols=4, style="Table Grid")

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '진행상태'
    hdr_cells[1].text = '리전'
    hdr_cells[2].text = '제목'
    hdr_cells[3].text = '작업내용'    
    
    hdr_cells[0].width = Cm(0.05)
    hdr_cells[1].width = Cm(0.05)
    hdr_cells[2].width = Cm(0.1)
    hdr_cells[3].width = Cm(50)
    
    for ind in df.index:
        row_cells = table.add_row().cells
        row_cells[0].text = df['진행상태'][ind]
        row_cells[1].text = df['리전'][ind]
        row_cells[2].text = df['제목'][ind]
        # row_cells[3].text = df['작업내용'][ind]  
        
        cell_content = df['작업내용'][ind]
        convert_markdown_to_rtf(row_cells[3], cell_content)
        
        for cell in row_cells:
            set_cell_font(cell, font_size=10)
            
def addTable4(df):
    
    # if [ month != '누적' ]:
    #     df = df[df['월'] == month].reset_index()
    # else:
    #     pass
    
    table = document.add_table(rows=1, cols=3, style="Table Grid")

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '진행상태'
    hdr_cells[1].text = '제목'
    hdr_cells[2].text = '작업내용'    

    for ind in df.index:
        row_cells = table.add_row().cells
        row_cells[0].text = df['진행상태'][ind]
        row_cells[1].text = df['제목'][ind]
        # row_cells[2].text = df['작업내용'][ind]  
        
        cell_content = df['작업내용'][ind]
        convert_markdown_to_rtf(row_cells[2], cell_content)
        
        for cell in row_cells:
            set_cell_font(cell, font_size=10)
            
def addTable_issue_all(df):
    
    # if [ month != '누적' ]:
    #     df = df[df['월'] == month].reset_index()
    # else:
    #     pass
    
    df.sort_values(by=['진행상태','리전','AZ'], axis=0, ascending=True, inplace=True)
    table = document.add_table(rows=1, cols=4, style="Table Grid")
    table.allow_autofit = True
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '진행상태'
    hdr_cells[1].text = '리전'
    hdr_cells[2].text = 'AZ'
    hdr_cells[3].text = '제목'
    
    hdr_cells[0].width = Cm(0.05)
    hdr_cells[1].width = Cm(0.05)
    hdr_cells[2].width = Cm(0.05)
    hdr_cells[3].width = Cm(50)

    for ind in df.index:
        row_cells = table.add_row().cells
        row_cells[0].text = df['진행상태'][ind]
        row_cells[1].text = df['리전'][ind]
        row_cells[2].text = str(df['AZ'][ind])
        row_cells[3].text = df['제목'][ind]
        for cell in row_cells:
            set_cell_font(cell, font_size=10)
            
def addTable_issue(df):
    
#     if [ month != '누적' ]:
#         df = df[df['월'] == month].reset_index()
#     else:
#         pass
    
    table = document.add_table(rows=1, cols=2, style="Table Grid")

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '진행상태'
    hdr_cells[1].text = '제목'

    for ind in df.index:
        row_cells = table.add_row().cells
        row_cells[0].text = df['진행상태'][ind]
        row_cells[1].text = df['제목'][ind]
        for cell in row_cells:
            set_cell_font(cell, font_size=10)

def addPivot2Table(df):
    
    table = document.add_table(rows=1, cols=9, style="Table Grid")
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '유형'
    hdr_cells[1].text = 'KR'
    hdr_cells[2].text = 'NA'
    hdr_cells[3].text = 'EU'    
    hdr_cells[4].text = 'CN'
    hdr_cells[5].text = 'SG'
    hdr_cells[6].text = 'RU'
    hdr_cells[7].text = '합계' 
    hdr_cells[8].text = '비중' 
    
    for ind in df.index:
        row_cells = table.add_row().cells
        row_cells[0].text = df['유형'][ind]
        row_cells[1].text = '%0.1f' % df['KR'][ind]
        row_cells[2].text = '%0.1f' % df['NA'][ind]
        row_cells[3].text = '%0.1f' % df['EU'][ind]  
        row_cells[4].text = '%0.1f' % df['CN'][ind]
        row_cells[5].text = '%0.1f' % df['SG'][ind]
        row_cells[6].text = '%0.1f' % df['RU'][ind]
        row_cells[7].text = '%0.1f' % df['합계'][ind]  
        row_cells[8].text = '%0.1f' % df['비중'][ind]  
        for cell in row_cells:
            set_cell_font(cell, font_size=10)
            
def addFailedTable(df,  month):
    
    if [ month != '누적' ]:
        df = df[df['월'] == month].reset_index()
    else:
        pass
    
    table = document.add_table(rows=1, cols=5, style="Table Grid")

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '날짜'
    hdr_cells[1].text = '진행상태'
    hdr_cells[2].text = '장애이벤트'    
    hdr_cells[3].text = '조치내용'
    hdr_cells[4].text = '장애인지소요시간(분)'
    
    for ind in df.index:
        row_cells = table.add_row().cells
        row_cells[0].text = df['장애발생시간'][ind]
        row_cells[1].text = df['진행상태'][ind]
        row_cells[2].text = df['장애이벤트'][ind]
        row_cells[3].text = df['조치내용'][ind]  
        row_cells[4].text = '%0.1f' % df['장애인지소요시간'][ind]
        for cell in row_cells:
            set_cell_font(cell, font_size=10)
        
def addTable_DB_license(df,  month):
    
    if [ month != '누적' ]:
        df = df[df['월'] == month].reset_index()
    else:
        pass
    
    table = document.add_table(rows=1, cols=7, style="Table Grid")

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '년'
    hdr_cells[1].text = '월'
    hdr_cells[2].text = '리전'
    hdr_cells[3].text = '라이선스 종류'    
    hdr_cells[4].text = 'Hostname'    
    hdr_cells[5].text = '라이선스'    
    hdr_cells[6].text = '만료일'    

    for ind in df.index:
        row_cells = table.add_row().cells
        row_cells[0].text = str(df['년'][ind])
        row_cells[1].text = str(df['월'][ind])
        row_cells[2].text = df['리전'][ind]
        row_cells[3].text = df['유형'][ind]
        row_cells[4].text = df['instance_name'][ind]          
        row_cells[5].text = df['공급사'][ind]          
        row_cells[6].text = str(df['expired_at'][ind])
        for cell in row_cells:
            set_cell_font(cell, font_size=10)
            
def addTable(df):
    print("in")   
    table = document.add_table(df.shape[0]+1, df.shape[1], style="Table Grid")

    # add the header rows
    for j in range(df.shape[-1]):
        table.cell(0,j).text = df.columns[j]
        print("1번 반복문", j)

    # add the rest of the dataframe    
    for i in range(df.shape[0]):
        for j in range(df.shape[-1]):
            table.cell(i+1,j).text = str(df.values[i,j])
            print("2번 반복문 i : ",i, "j : ", j)
    print("finish")
    
    
def addTable_1_8(df_filtered,  month):
    
    # if [ month != '누적' ]:
    #     df = df[df['월'] == month].reset_index()
    # else:
    #     pass

    df_filtered = df_filtered[[ '리전', '유형', '성패', 'Instance', '비고']]
    df_filtered.sort_values(by=['비고', '리전'], axis=0, ascending=False, inplace=True)
    
    table = document.add_table(rows=1, cols=5, style="Table Grid")
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '리전'
    hdr_cells[1].text = '유형'
    hdr_cells[2].text = '성패'
    hdr_cells[3].text = 'Instance'
    hdr_cells[4].text = '비고'
    
    for ind in df_filtered.index:
        row_cells = table.add_row().cells
        # row_cells[0].text = str(df_filtered['월'][ind])
        row_cells[0].text = df_filtered['리전'][ind]
        row_cells[1].text = df_filtered['유형'][ind]
        row_cells[2].text = df_filtered['성패'][ind]
        row_cells[3].text = df_filtered['Instance'][ind]
        row_cells[4].text = str(df_filtered['비고'][ind])
        for cell in row_cells:
            set_cell_font(cell, font_size=10)
            
def addTable_1_8_all(df_filtered,  month):
    
    # if [ month != '누적' ]:
    #     df = df[df['월'] == month].reset_index()
    # else:
    #     pass

    # df_filtered = df_filtered[[ '연도', '월', '비고', '리전', 'AZ', '유형', 'Instance']]
    # df_filtered.sort_values(by=['리전'], axis=0, ascending=True, inplace=True)    
    # unique_df = df_filtered['비고'].value_counts()
    # unique_df = pd.DataFrame({'비고':unique_df.index, 'DB 인스턴스 Qty':unique_df.values})
    # unique_df.sort_values(by=['비고'], axis=0, ascending=True, inplace=True)
    # unique_df = unique_df.reset_index()
    # df_filtered = df_filtered.drop_duplicates(['비고'])
    
    df_filtered = df_filtered[[ '년', '월', '비고', '리전', 'AZ', '유형', 'DB 인스턴스 Qty']]
    df_filtered.sort_values(by=['비고', '리전'], axis=0, ascending=False, inplace=True)
    
    table = document.add_table(rows=1, cols=7, style="Table Grid")
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '년'
    hdr_cells[1].text = '월'
    hdr_cells[2].text = '미진행 사유'
    hdr_cells[3].text = '리전'
    hdr_cells[4].text = 'AZ'
    hdr_cells[5].text = '유형'
    hdr_cells[6].text = 'DB 인스턴스 Qty'
    
    # hdr_cells[0].width = Cm(0.05)
    # hdr_cells[1].width = Cm(0.05)
    # hdr_cells[2].width = Cm(50)
    # hdr_cells[3].width = Cm(0.05)
    # hdr_cells[4].width = Cm(0.05)
    # hdr_cells[5].width = Cm(1)
    # hdr_cells[6].width = Cm(10)
    
    for ind in df_filtered.index:
        row_cells = table.add_row().cells
        row_cells[0].text = str(df_filtered['년'][ind])
        row_cells[1].text = str(df_filtered['월'][ind])
        row_cells[2].text = df_filtered['비고'][ind]
        row_cells[3].text = df_filtered['리전'][ind]
        row_cells[4].text = str(df_filtered['AZ'][ind])
        row_cells[5].text = df_filtered['유형'][ind]
        row_cells[6].text = str(df_filtered['DB 인스턴스 Qty'][ind])
        for cell in row_cells:
            set_cell_font(cell, font_size=10)
        
def addFailedTable2(df,  month):
    
    if [ month != '누적' ]:
        df = df[df['월'] == month].reset_index()
    else:
        pass
    
    table = document.add_table(rows=1, cols=5, style="Table Grid")

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '날짜'
    hdr_cells[1].text = '진행상태'
    hdr_cells[2].text = '장애이벤트'    
    hdr_cells[3].text = '조치내용'
    hdr_cells[4].text = '장애전파소요시간(분)'
    
    for ind in df.index:
        row_cells = table.add_row().cells
        row_cells[0].text = df['장애발생시간'][ind]
        row_cells[1].text = df['진행상태'][ind]
        row_cells[2].text = df['장애이벤트'][ind]
        row_cells[3].text = df['조치내용'][ind]  
        row_cells[4].text = '%0.1f' % df['장애전파소요시간'][ind]
        for cell in row_cells:
            set_cell_font(cell, font_size=10)